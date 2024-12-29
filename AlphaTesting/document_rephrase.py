# Install better_profanity, pdfplumber, and python-docx
# !pip install better-profanity pdfplumber python-docx

# # Install NLTK for natural language processing (WordNet)
# !pip install nltk

# Download WordNet data for synonyms
import nltk
nltk.download('wordnet')
nltk.download('omw-1.4') 

import os
from better_profanity import profanity
from nltk.corpus import wordnet
from pdfplumber import open as open_pdf
from docx import Document
import re

# Initialize profanity filter
profanity.load_censor_words()

# Function to find synonyms or legal alternatives based on context
def get_legal_alternatives(word):
    """
    Finds legal alternatives to a profane word based on its synonyms in WordNet.
    """
    alternatives = []
    for syn in wordnet.synsets(word):
        for lemma in syn.lemmas():
            if lemma.name() != word and not profanity.contains_profanity(lemma.name()):
                alternatives.append(lemma.name().replace('_', ' '))  # Replace underscores with spaces
    return alternatives if alternatives else [word]  # Return the original word if no alternatives are found

# Function to replace profane words within the context of a sentence
def replace_profanity_in_sentence(sentence):
    """
    Replaces profane words in a sentence with legal alternatives.
    """
    words = sentence.split()
    replaced_sentence = []

    for word in words:
        # Remove punctuation for proper profanity comparison
        cleaned_word = re.sub(r'[^\w\s]', '', word)
        if profanity.contains_profanity(cleaned_word):
            alternatives = get_legal_alternatives(cleaned_word)
            # Replace with the first legal alternative found
            replaced_sentence.append(alternatives[0])
        else:
            replaced_sentence.append(word)

    return ' '.join(replaced_sentence)

# Function to replace profanity across an entire text based on sentence context
def replace_profanity_in_text(text):
    """
    Processes text, replacing any profane words based on sentence-level context.
    """
    # Split the text into sentences for context-aware processing
    sentences = re.split(r'(?<=[.!?]) +', text)
    replaced_sentences = [replace_profanity_in_sentence(sentence) for sentence in sentences]
    return ' '.join(replaced_sentences)

# Function to check and replace profanity
def check_and_replace_profanity(text):
    """
    Detects and replaces profane words in the text, returning whether the text contained profanity.
    """
    contains_profanity = profanity.contains_profanity(text)
    replaced_text = replace_profanity_in_text(text)
    return contains_profanity, replaced_text

# Function to read text from a file (supports .txt, .pdf, .docx)
def read_file(file_path):
    """
    Reads the content from a file and returns the text.
    Supports: .txt, .pdf, .docx.
    """
    _, file_extension = os.path.splitext(file_path)

    if file_extension == '.txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    elif file_extension == '.pdf':
        with open_pdf(file_path) as pdf:
            return ' '.join(page.extract_text() for page in pdf.pages if page.extract_text())
    elif file_extension == '.docx':
        doc = Document(file_path)
        return '\n'.join(paragraph.text for paragraph in doc.paragraphs)
    else:
        raise ValueError(f"Unsupported file format: {file_extension}")

# Function to write the output to a file, preserving format
def write_output(input_file_path, output_text):
    """
    Writes the censored output to the same file format as the original.
    Supports: .txt, .pdf (output warning), .docx.
    """
    _, file_extension = os.path.splitext(input_file_path)
    output_file_path = f"censored_output{file_extension}"

    if file_extension == '.txt':
        with open(output_file_path, 'w', encoding='utf-8') as f:
            f.write(output_text)
    elif file_extension == '.pdf':
        print("PDF output not supported yet. Please save manually.")
    elif file_extension == '.docx':
        doc = Document()
        for paragraph in output_text.split('\n'):
            doc.add_paragraph(paragraph)
        doc.save(output_file_path)
    else:
        raise ValueError(f"Unsupported file format: {file_extension}")

# Main function to process the input file
def process_file(input_file_path):
    """
    Processes the input file for profanity and writes the censored output in the same format.
    """
    # Read input file
    input_text = read_file(input_file_path)

    # Check for and replace profanity
    contains_profanity, replaced_text = check_and_replace_profanity(input_text)

    # Show whether the input contained profanity and provide output details
    print(f'Input Text: "{input_text[:100]}..." - Contains Profanity: {"Yes" if contains_profanity else "No"}')
    
    # Write the replaced text to an output file
    write_output(input_file_path, replaced_text)
    print(f'Censored output saved as: censored_output{os.path.splitext(input_file_path)[1]}')

# Example usage (input file should be present in Colab environment or your local system)
input_file_path = '/content/Transcript.txt'  # Change this path as per your file location
process_file(input_file_path)
