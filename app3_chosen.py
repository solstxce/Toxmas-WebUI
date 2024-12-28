import os
import re
import streamlit as st
from io import BytesIO
from docx import Document
from PyPDF2 import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import Paragraph, Frame
from reportlab.lib.units import inch
from better_profanity import profanity
import pdfplumber
import datetime

# Initialize profanity filter
profanity.load_censor_words()

# Function to find synonyms or legal alternatives based on context
def get_legal_alternatives(word):
    from nltk.corpus import wordnet
    alternatives = []
    for syn in wordnet.synsets(word):
        for lemma in syn.lemmas():
            if lemma.name() != word and not profanity.contains_profanity(lemma.name()):
                alternatives.append(lemma.name().replace('_', ' '))
    return alternatives if alternatives else [word]

# Function to replace profane words within the context of a sentence
def replace_profanity_in_sentence(sentence):
    words = sentence.split()
    replaced_sentence = []

    for word in words:
        cleaned_word = re.sub(r'[^\w\s]', '', word)
        if profanity.contains_profanity(cleaned_word):
            alternatives = get_legal_alternatives(cleaned_word)
            replaced_sentence.append(alternatives[0])
        else:
            replaced_sentence.append(word)

    return ' '.join(replaced_sentence)

# Function to replace profanity across an entire text based on sentence context
def replace_profanity_in_text(text):
    sentences = re.split(r'(?<=[.!?]) +', text)
    replaced_sentences = [replace_profanity_in_sentence(sentence) for sentence in sentences]
    return ' '.join(replaced_sentences)

# Function to check and replace profanity
def check_and_replace_profanity(text):
    contains_profanity = profanity.contains_profanity(text)
    replaced_text = replace_profanity_in_text(text)
    return contains_profanity, replaced_text

# Function to check profanity in the text and return a list of profane words
def check_profanity(text):
    contains_profanity = profanity.contains_profanity(text)
    return "Profane" if contains_profanity else "Not Profane"

# Function to read text from a file (supports .txt, .pdf, .docx)
def read_file(file_path):
    _, file_extension = os.path.splitext(file_path)

    if file_extension == '.txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    elif file_extension == '.pdf':
        with pdfplumber.open(file_path) as pdf:
            return ' '.join(page.extract_text() for page in pdf.pages if page.extract_text())
    elif file_extension == '.docx':
        doc = Document(file_path)
        paragraphs = [paragraph.text for paragraph in doc.paragraphs]
        return '\n\n'.join(paragraphs)
    else:
        raise ValueError(f"Unsupported file format: {file_extension}")

# Function to analyze document for profanity and generate a PDF report
def analyze_document(input_path):
    file_extension = os.path.splitext(input_path)[1].lower()
    input_filename = os.path.basename(input_path)
    output_filename = f"Report for {os.path.splitext(input_filename)[0]}.pdf"
    output_path = os.path.join(os.path.dirname(input_path), output_filename)

    if file_extension == '.txt':
        with open(input_path, 'r', encoding='utf-8') as file:
            content = file.read()
            paragraphs = content.split('\n\n')
    elif file_extension == '.docx':
        doc = Document(input_path)
        paragraphs = [paragraph.text for paragraph in doc.paragraphs]
    elif file_extension == '.pdf':
        reader = PdfReader(input_path)
        paragraphs = []
        for page in reader.pages:
            paragraphs.extend(page.extract_text().split('\n\n'))
    else:
        raise ValueError("Unsupported file format")

    words = ' '.join(paragraphs).split()
    total_words = len(words)

    profane_words = [word for word in words if check_profanity(word) == "Profane"]
    profane_word_count = len(profane_words)

    create_clean_pdf_output(output_path, paragraphs, profane_words)
    add_watermark(output_path, "Verified by TOX-MAS")
    return output_path

# Function to create PDF output
def create_clean_pdf_output(output_path, paragraphs, profane_words):
    packet = BytesIO()
    can = canvas.Canvas(packet, pagesize=letter)
    width, height = letter

    y = height - inch  # Start writing just below the top margin
    for paragraph in paragraphs:
        words = paragraph.split()
        x = inch
        for word in words:
            is_profane = word in profane_words

            if x + can.stringWidth(word + " ") > width - inch:
                y -= 20
                x = inch
                if y < inch:
                    can.showPage()
                    y = height - inch

            if is_profane:
                # Censor profane words (optional: can replace with '[CENSORED]' if needed)
                can.setFillColor(colors.black)
                can.rect(x, y - 2, can.stringWidth(word), 14, fill=1)
            else:
                can.setFont("Helvetica", 12)
                can.setFillColor(colors.black)
                can.drawString(x, y, word)

            x += can.stringWidth(word + " ")

        y -= 20  # Add space between paragraphs

    can.save()

    packet.seek(0)
    new_pdf = PdfReader(packet)
    output = PdfWriter()

    for page in new_pdf.pages:
        output.add_page(page)

    with open(output_path, "wb") as output_stream:
        output.write(output_stream)

from reportlab.lib.colors import green

# Function to add watermark (green tick and text) to each page of the PDF
def add_watermark(input_pdf_path, watermark_text="Verified by TOX-MAS"):
    try:
        output_pdf_path = input_pdf_path.replace(".pdf", "_watermarked.pdf")
        packet = BytesIO()
        can = canvas.Canvas(packet, pagesize=letter)
        width, height = letter

        # Add watermark to each page (a green tick + "Verified by TOX-MAS")
        can.setFillColor(colors.black)  # Set the fill color to black for the tick
        can.setFont("Helvetica", 40)
        tick = u'\u2713'  # Unicode for green tick symbol

        for _ in range(1):  # Since this is a packet canvas, create only one page (base watermark)
            # Draw tick at the bottom right
            #can.drawString(width - 2*inch, inch, tick)
            # Draw the "Verified by TOX-MAS" text below the tick
            can.setFont("Helvetica", 10)
            can.drawString(width - 2*inch, inch - 20, watermark_text)
            # Optionally, you can rotate and place the text diagonally across the page
            #can.setFillColor(black)  # Set text color to green
            can.saveState()
            #can.translate(width / 2, height / 2)
            #can.rotate(45)
            #can.setFont("Helvetica", 20)
            #can.drawCentredString(0, 0, watermark_text)
            can.restoreState()
            can.showPage()

        can.save()

        packet.seek(0)
        watermark_pdf = PdfReader(packet)
        input_pdf = PdfReader(open(input_pdf_path, "rb"))
        output_pdf = PdfWriter()

        # Merge the watermark onto each page of the input PDF
        for i in range(len(input_pdf.pages)):
            page = input_pdf.pages[i]
            page.merge_page(watermark_pdf.pages[0])  # Merge the watermark page onto the current page
            output_pdf.add_page(page)

        # Save the final watermarked PDF
        with open(output_pdf_path, "wb") as output_stream:
            output_pdf.write(output_stream)

        # Replace the original file with the watermarked version
        if os.path.exists(input_pdf_path) and input_pdf_path.endswith('.pdf'):
            input_pdf.stream.close()  # Ensure the file is closed before removing
            os.remove(input_pdf_path)
            os.rename(output_pdf_path, input_pdf_path)
        else:
            raise FileNotFoundError(f"{input_pdf_path} is not found or not a valid PDF.")
            
    except Exception as e:
        print(f"An error occurred while adding the watermark: {e}")

# Function to log past analysis
def log_past_analysis(file_name, mode, profane_word_count):
    if "analysis_log" not in st.session_state:
        st.session_state["analysis_log"] = []
    st.session_state["analysis_log"].append({
        "file_name": file_name,
        "mode": mode,
        "profane_word_count": profane_word_count,
        "timestamp": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    })

# Streamlit UI
st.title("Profanity Analyzer ")

uploaded_file = st.file_uploader("Upload a document", type=["txt", "pdf", "docx"])

option = st.selectbox("Choose an option", ["Choose...", "Censor", "Rephrase"])

if st.button("Analyze"):
    if uploaded_file:
        file_details = {"filename": uploaded_file.name, "filetype": uploaded_file.type, "filesize": uploaded_file.size}
        st.write(file_details)

        input_path = os.path.join("temp", uploaded_file.name)
        with open(input_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        if option == "Censor":
            output_path = analyze_document(input_path)
            profane_word_count = len([word for word in read_file(input_path).split() if check_profanity(word) == "Profane"])
            total_words = len(read_file(input_path).split())
            st.success("Document analyzed and censored successfully.")
            st.write(f"Total words: {total_words}")
            st.write(f"Profane words: {profane_word_count}")
            st.write(f"Percentage of profane words: {(profane_word_count / total_words) * 100:.2f}%")
            with open(output_path, "rb") as file:
                st.download_button(label="Download Censored Document", data=file, file_name=output_path)
            log_past_analysis(uploaded_file.name, "Censor", profane_word_count)

        elif option == "Rephrase":
            input_paragraphs = read_file(input_path)
            output_paragraphs = []
            contains_profanity = False

            for paragraph in input_paragraphs.split('\n\n'):
                paragraph_contains_profanity, replaced_paragraph = check_and_replace_profanity(paragraph)
                contains_profanity |= paragraph_contains_profanity
                output_paragraphs.append(replaced_paragraph)

            output_file_path = f"rephrased_output.pdf"
            create_clean_pdf_output(output_file_path, output_paragraphs, [])
            add_watermark(output_file_path, "Verified by TOX-MAS")
            #total_words = len(' '.join(output_paragraphs).split())
            profane_word_count = len([word for word in read_file(input_path).split() if check_profanity(word) == "Profane"])
            total_words = len(read_file(input_path).split())
            st.success("Document analyzed and rephrased successfully.")
            st.write(f"Total words: {total_words}")
            st.write(f"Percentage of profane words: {(profane_word_count / total_words) * 100:.2f}%")
            st.write(f"Number of Profane words rephrased are: {profane_word_count}")
            with open(output_file_path, "rb") as file:
                st.download_button(label="Download Rephrased Document", data=file, file_name=output_file_path)
            log_past_analysis(uploaded_file.name, "Rephrase", profane_word_count)
            
# Display past analysis
st.header("Past Analysis")
if "analysis_log" in st.session_state and len(st.session_state["analysis_log"]) > 0:
    for entry in st.session_state["analysis_log"]:
        st.write(f"- **File:** {entry['file_name']}, **Mode:** {entry['mode']}, **Profane words:** {entry['profane_word_count']}, **Date:** {entry['timestamp']}")
else:
    st.write("No analysis records found.")